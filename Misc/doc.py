import docx

# Create an instance to the word document


document = docx.Document()
# Adding heading
document.add_heading('Python Document')
# Adding paragraphs
document.add_paragraph('''This is a document generated using Python script. It is done using python docx module.''')
document.add_paragraph('''Python supports the use of script to perform word processing actions.''')

#Saving document

document.save('sample.docx')

# Reading a docx file.
d1 = docx.Document('sample.docx')

for para in d1.paragraphs:
    print(para.text)